"""Unit tests for parsers and OCR utilities."""
import io
import pytest
from PIL import Image


# ── Image parser ──────────────────────────────────────────────────────────────

def _make_jpeg(width: int = 400, height: int = 300, color=(120, 180, 60)) -> bytes:
    img = Image.new("RGB", (width, height), color=color)
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    return buf.getvalue()


def _make_png(width: int = 200, height: int = 800, color=(255, 255, 255)) -> bytes:
    img = Image.new("RGB", (width, height), color=color)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def test_image_parser_jpeg():
    from workers.ingestion_worker.parsers.image_parser import parse_image

    result = parse_image(_make_jpeg(640, 480))
    assert result.width == 640
    assert result.height == 480
    assert len(result.thumbnail_bytes) > 0


def test_image_parser_thumbnail_is_smaller():
    from workers.ingestion_worker.parsers.image_parser import parse_image, THUMBNAIL_SIZE

    result = parse_image(_make_jpeg(2000, 2000))
    thumb = Image.open(io.BytesIO(result.thumbnail_bytes))
    assert thumb.width <= THUMBNAIL_SIZE[0]
    assert thumb.height <= THUMBNAIL_SIZE[1]


def test_image_parser_text_likely_for_tall_image():
    """A tall portrait image (document-like) should be flagged as text_likely."""
    from workers.ingestion_worker.parsers.image_parser import parse_image

    result = parse_image(_make_png(400, 1200))
    assert result.text_likely is True


def test_image_parser_text_not_likely_for_wide_panorama():
    """A wide panorama (aspect ratio > 2) should not be flagged as text_likely."""
    from workers.ingestion_worker.parsers.image_parser import parse_image

    result = parse_image(_make_jpeg(2400, 400))
    assert result.text_likely is False


# ── OCR utilities ─────────────────────────────────────────────────────────────

def test_ocr_needed_empty():
    from workers.ingestion_worker.ocr.tesseract_ocr import ocr_needed
    assert ocr_needed("") is True


def test_ocr_needed_sparse():
    from workers.ingestion_worker.ocr.tesseract_ocr import ocr_needed
    assert ocr_needed("x" * 49) is True


def test_ocr_not_needed_sufficient_text():
    from workers.ingestion_worker.ocr.tesseract_ocr import ocr_needed
    assert ocr_needed("x" * 50) is False
    # 51 chars — clearly above threshold
    assert ocr_needed("This is a full page of text with enough content here.") is False


# ── DOCX parser ───────────────────────────────────────────────────────────────

def _make_minimal_docx() -> bytes:
    """Create a minimal valid DOCX file in memory."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph.")
    doc.add_paragraph("Second paragraph with more content.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_parser_extracts_blocks():
    from workers.ingestion_worker.parsers.docx_parser import parse_docx

    docx_bytes = _make_minimal_docx()
    result = parse_docx(docx_bytes)
    assert len(result.blocks) >= 3  # heading + 2 paragraphs
    contents = [b.content for b in result.blocks]
    assert "Test Heading" in contents
    assert "This is a test paragraph." in contents


def test_docx_parser_identifies_heading_type():
    from workers.ingestion_worker.parsers.docx_parser import parse_docx

    result = parse_docx(_make_minimal_docx())
    types = [b.block_type for b in result.blocks]
    assert "heading" in types


# ── PDF parser ────────────────────────────────────────────────────────────────

def _make_minimal_pdf() -> bytes:
    """Create a minimal text-based PDF with PyMuPDF."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from page one. " * 20)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pdf_parser_extracts_page():
    from workers.ingestion_worker.parsers.pdf_parser import parse_pdf

    result = parse_pdf(_make_minimal_pdf())
    assert result.page_count == 1
    assert len(result.pages) == 1
    assert "Hello from page one" in result.pages[0].text


def test_pdf_parser_marks_scanned_pages_for_ocr():
    """A page with very little text should be flagged needs_ocr=True."""
    import fitz

    doc = fitz.open()
    doc.new_page()  # empty page — no text
    buf = io.BytesIO()
    doc.save(buf)

    from workers.ingestion_worker.parsers.pdf_parser import parse_pdf

    result = parse_pdf(buf.getvalue())
    assert result.pages[0].needs_ocr is True


def test_pdf_parser_respects_max_pages():
    import fitz

    doc = fitz.open()
    for i in range(5):
        page = doc.new_page()
        page.insert_text((72, 72), f"Page {i}")
    buf = io.BytesIO()
    doc.save(buf)

    from workers.ingestion_worker.parsers.pdf_parser import parse_pdf

    result = parse_pdf(buf.getvalue(), max_pages=3)
    assert result.page_count == 3
