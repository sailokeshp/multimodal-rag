"""
DOCX parser: headings, paragraphs, bullet points, tables, inline images.
Requires: python-docx
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocxBlock:
    block_type: str   # heading | paragraph | list_item | table_cell | caption
    section_title: str | None
    content: str


@dataclass
class DocxParseResult:
    blocks: list[DocxBlock]
    inline_image_bytes: list[bytes]


def parse_docx(file_bytes: bytes) -> DocxParseResult:
    """
    Extract structured blocks and inline images from a DOCX file.
    """
    from docx import Document
    from docx.oxml.ns import qn
    import zipfile
    import io

    doc = Document(io.BytesIO(file_bytes))
    blocks: list[DocxBlock] = []
    current_section: str | None = None

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style.startswith("Heading"):
            current_section = text
            blocks.append(DocxBlock("heading", current_section, text))
        elif style.startswith("List"):
            blocks.append(DocxBlock("list_item", current_section, text))
        else:
            blocks.append(DocxBlock("paragraph", current_section, text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    blocks.append(DocxBlock("table_cell", current_section, cell_text))

    # Extract inline images from the zip archive
    inline_images: list[bytes] = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    inline_images.append(z.read(name))
    except Exception as exc:
        logger.warning("Could not extract inline images from DOCX: %s", exc)

    logger.info("DOCX parsed: blocks=%d inline_images=%d", len(blocks), len(inline_images))
    return DocxParseResult(blocks=blocks, inline_image_bytes=inline_images)
